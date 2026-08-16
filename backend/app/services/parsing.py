"""Document parsing.

Produces `ParsedBlock`s that keep structure: heading level, page number, and
whether the block is text or an image. Chunking consumes these and never splits
across a heading boundary without carrying the heading trail forward.
"""

from __future__ import annotations

import csv
import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

TEXT_MIMES = {
    "text/plain", "text/markdown", "text/x-markdown", "application/json",
    "text/csv", "text/html", "application/xml", "text/xml",
}
IMAGE_MIMES = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif", "image/bmp"}
PDF_MIMES = {"application/pdf"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class UnsupportedFileType(ValueError):
    pass


@dataclass(slots=True)
class ParsedBlock:
    text: str
    page: int | None = None
    heading_path: list[str] = field(default_factory=list)
    kind: str = "text"          # text | heading | table | image
    image_uri: str | None = None


@dataclass(slots=True)
class ParsedDocument:
    blocks: list[ParsedBlock]
    page_count: int = 0
    metadata: dict = field(default_factory=dict)


def parse(path: str, mime_type: str, *, storage_uri: str | None = None) -> ParsedDocument:
    mime = (mime_type or "").split(";")[0].strip().lower()
    suffix = Path(path).suffix.lower()

    if mime in PDF_MIMES or suffix == ".pdf":
        return _parse_pdf(path)
    if mime in DOCX_MIMES or suffix == ".docx":
        return _parse_docx(path)
    if mime in IMAGE_MIMES or suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}:
        return _parse_image(path, storage_uri or f"file://{path}")
    if mime == "text/csv" or suffix == ".csv":
        return _parse_csv(path)
    if mime in {"text/markdown", "text/x-markdown"} or suffix in {".md", ".markdown"}:
        return _parse_markdown(_read_text(path))
    if mime in {"text/html", "application/xhtml+xml"} or suffix in {".html", ".htm"}:
        return _parse_html(_read_text(path))
    if mime in TEXT_MIMES or suffix in {".txt", ".json", ".log", ".yaml", ".yml", ".py", ".ts", ".tsx", ".js"}:
        return _parse_plain(_read_text(path))

    # Last resort: if it decodes as UTF-8 it is close enough to text.
    try:
        return _parse_plain(_read_text(path))
    except UnicodeDecodeError as exc:
        raise UnsupportedFileType(f"Cannot parse {mime or suffix or 'file'}") from exc


def _read_text(path: str) -> str:
    return Path(path).read_text(encoding="utf-8", errors="replace")


# ------------------------------------------------------------------- pdf ----
def _parse_pdf(path: str) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(path)
    blocks: list[ParsedBlock] = []
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 - a single bad page must not fail the doc
            logger.warning("pdf page %d extraction failed: %s", page_no, exc)
            continue
        text = _normalize_ws(text)
        if text:
            blocks.append(ParsedBlock(text=text, page=page_no))

    meta = {}
    try:
        info = reader.metadata or {}
        meta = {k.lstrip("/"): str(v) for k, v in info.items() if v}
    except Exception:  # noqa: BLE001
        pass

    if not blocks:
        logger.warning("no extractable text in %s — likely a scanned PDF needing OCR", path)
    return ParsedDocument(blocks=blocks, page_count=len(reader.pages), metadata=meta)


# ------------------------------------------------------------------ docx ----
def _parse_docx(path: str) -> ParsedDocument:
    import docx

    doc = docx.Document(path)
    blocks: list[ParsedBlock] = []
    heading_path: list[str] = []

    for para in doc.paragraphs:
        text = _normalize_ws(para.text)
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            level = _heading_level(style)
            heading_path = heading_path[: level - 1] + [text]
            blocks.append(ParsedBlock(text=text, kind="heading", heading_path=list(heading_path)))
        else:
            blocks.append(ParsedBlock(text=text, heading_path=list(heading_path)))

    for table in doc.tables:
        rendered = _render_table([[c.text for c in row.cells] for row in table.rows])
        if rendered:
            blocks.append(ParsedBlock(text=rendered, kind="table", heading_path=list(heading_path)))

    return ParsedDocument(blocks=blocks, page_count=0)


def _heading_level(style: str) -> int:
    m = re.search(r"(\d+)", style)
    return max(1, min(int(m.group(1)), 6)) if m else 1


# -------------------------------------------------------------- markdown ----
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def _parse_markdown(text: str) -> ParsedDocument:
    blocks: list[ParsedBlock] = []
    heading_path: list[str] = []
    buffer: list[str] = []
    in_code = False

    def flush() -> None:
        joined = "\n".join(buffer).strip()
        if joined:
            blocks.append(ParsedBlock(text=joined, heading_path=list(heading_path)))
        buffer.clear()

    for line in text.splitlines():
        if line.lstrip().startswith("```"):
            in_code = not in_code
            buffer.append(line)
            continue
        m = _MD_HEADING.match(line) if not in_code else None
        if m:
            flush()
            level, title = len(m.group(1)), m.group(2).strip()
            heading_path = heading_path[: level - 1] + [title]
            blocks.append(ParsedBlock(text=title, kind="heading", heading_path=list(heading_path)))
        else:
            buffer.append(line)
    flush()
    return ParsedDocument(blocks=blocks)


# ------------------------------------------------------------------ html ----
def _parse_html(html: str) -> ParsedDocument:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return _parse_plain(re.sub(r"<[^>]+>", " ", html))

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer"]):
        tag.decompose()

    blocks: list[ParsedBlock] = []
    heading_path: list[str] = []
    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "td"]):
        text = _normalize_ws(el.get_text(" "))
        if not text:
            continue
        if el.name.startswith("h") and len(el.name) == 2 and el.name[1].isdigit():
            level = int(el.name[1])
            heading_path = heading_path[: level - 1] + [text]
            blocks.append(ParsedBlock(text=text, kind="heading", heading_path=list(heading_path)))
        else:
            blocks.append(ParsedBlock(text=text, heading_path=list(heading_path)))
    return ParsedDocument(blocks=blocks)


# ------------------------------------------------------------------- csv ----
def _parse_csv(path: str, max_rows_per_block: int = 40) -> ParsedDocument:
    raw = _read_text(path)
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        return ParsedDocument(blocks=[])

    header, body = rows[0], rows[1:]
    blocks: list[ParsedBlock] = []
    # Repeat the header in every block so a chunk is self-describing after retrieval.
    for i in range(0, len(body), max_rows_per_block):
        window = body[i : i + max_rows_per_block]
        blocks.append(
            ParsedBlock(
                text=_render_table([header, *window]),
                kind="table",
                heading_path=[Path(path).name],
            )
        )
    return ParsedDocument(blocks=blocks, metadata={"rows": len(body), "columns": len(header)})


def _render_table(rows: list[list[str]]) -> str:
    cleaned = [[_normalize_ws(c) for c in row] for row in rows if any(c.strip() for c in row)]
    return "\n".join(" | ".join(row) for row in cleaned)


# ----------------------------------------------------------------- image ----
def _parse_image(path: str, storage_uri: str) -> ParsedDocument:
    """Images become a single image-modality block. The multimodal embedding
    model vectorises the picture directly — no captioning step in between."""
    name = Path(path).name
    return ParsedDocument(
        blocks=[ParsedBlock(text=name, kind="image", image_uri=storage_uri, heading_path=[name])],
        page_count=1,
    )


# ----------------------------------------------------------------- plain ----
def _parse_plain(text: str) -> ParsedDocument:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    return ParsedDocument(blocks=[ParsedBlock(text=p) for p in paragraphs])


def _normalize_ws(text: str) -> str:
    text = text.replace("\xa0", " ").replace("​", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
